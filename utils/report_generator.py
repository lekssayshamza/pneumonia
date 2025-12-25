from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.platypus import Image as ReportLabImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from PIL import Image as PILImage
import io
from datetime import datetime
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import base64

def create_digital_signature():
    """
    Create a professional digital signature image for reports
    Returns a PIL Image with the signature
    """
    # Create signature image
    width, height = 400, 120
    img = PILImage.new('RGB', (width, height), color='white')
    
    from PIL import ImageDraw, ImageFont
    
    draw = ImageDraw.Draw(img)
    
    # Try to use a nice font, fallback to default if not available
    try:
        # Try to use a professional font
        font_large = ImageFont.truetype("arial.ttf", 28)
        font_small = ImageFont.truetype("arial.ttf", 12)
    except:
        try:
            font_large = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 28)
            font_small = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 12)
        except:
            font_large = ImageFont.load_default()
            font_small = ImageFont.load_default()
    
    # Draw medical cross icon (simple representation)
    icon_size = 30
    icon_x = 20
    icon_y = (height - icon_size) // 2
    
    # Draw a simple medical cross
    cross_color = '#2c5aa0'
    line_width = 4
    # Vertical line
    draw.rectangle([icon_x + icon_size//2 - line_width//2, icon_y, 
                    icon_x + icon_size//2 + line_width//2, icon_y + icon_size], 
                   fill=cross_color)
    # Horizontal line
    draw.rectangle([icon_x, icon_y + icon_size//2 - line_width//2, 
                    icon_x + icon_size, icon_y + icon_size//2 + line_width//2], 
                   fill=cross_color)
    
    # Draw organization name
    org_name = "Keep Awareness Alive"
    text_color = '#1a1a1a'
    
    # Get text dimensions
    bbox = draw.textbbox((0, 0), org_name, font=font_large)
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]
    
    # Position text next to icon
    text_x = icon_x + icon_size + 20
    text_y = (height - text_height) // 2 - 10
    
    draw.text((text_x, text_y), org_name, fill=text_color, font=font_large)
    
    # Draw tagline
    tagline = "Medical AI Analysis Platform"
    tagline_color = '#666666'
    bbox_tagline = draw.textbbox((0, 0), tagline, font=font_small)
    tagline_width = bbox_tagline[2] - bbox_tagline[0]
    tagline_height = bbox_tagline[3] - bbox_tagline[1]
    
    tagline_x = text_x
    tagline_y = text_y + text_height + 5
    
    draw.text((tagline_x, tagline_y), tagline, fill=tagline_color, font=font_small)
    
    # Draw decorative line
    line_y = height - 25
    draw.line([text_x, line_y, width - 20, line_y], fill='#cccccc', width=1)
    
    # Add date
    date_text = datetime.now().strftime("%B %d, %Y")
    date_bbox = draw.textbbox((0, 0), date_text, font=font_small)
    date_width = date_bbox[2] - date_bbox[0]
    date_x = width - date_width - 20
    date_y = line_y + 5
    
    draw.text((date_x, date_y), date_text, fill=tagline_color, font=font_small)
    
    return img

def create_pdf_report(username, email, prediction_label, confidence, original_image, heatmap_image, output_path, 
                      patient_name=None, patient_id=None, date_of_birth=None, gender=None, age=None):
    """
    Create a professional PDF report for the patient
    
    Args:
        username: Doctor username
        email: Doctor email
        prediction_label: "Normal" or "Pneumonia"
        confidence: Confidence score (0-1)
        original_image: PIL Image of original X-ray
        heatmap_image: PIL Image of heatmap overlay
        output_path: Path to save the PDF
        patient_name: Patient name (optional)
        patient_id: Patient ID (optional)
        date_of_birth: Patient date of birth (optional)
        gender: Patient gender (optional)
        age: Patient age (optional)
    """
    # Create PDF document
    doc = SimpleDocTemplate(output_path, pagesize=letter,
                          rightMargin=0.75*inch, leftMargin=0.75*inch,
                          topMargin=0.75*inch, bottomMargin=0.75*inch)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    normal_style = ParagraphStyle(
        'NormalText',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#333333'),
        spaceAfter=6,
        alignment=TA_LEFT
    )
    
    # Title
    title = Paragraph("Chest X-Ray Analysis Report", title_style)
    elements.append(title)
    elements.append(Spacer(1, 0.2*inch))
    
    # Patient Information Section
    elements.append(Paragraph("Patient Information", heading_style))
    
    # Create patient info table
    patient_data = []
    if patient_name:
        patient_data.append(['Patient Name:', patient_name])
    if patient_id:
        patient_data.append(['Patient ID:', patient_id])
    if date_of_birth:
        patient_data.append(['Date of Birth:', date_of_birth])
    if gender:
        patient_data.append(['Gender:', gender])
    if age:
        patient_data.append(['Age:', str(age)])
    
    # Doctor information
    patient_data.append(['Doctor Name:', username])
    patient_data.append(['Doctor Email:', email])
    patient_data.append(['Report Date:', datetime.now().strftime("%B %d, %Y at %I:%M %p")])
    
    patient_table = Table(patient_data, colWidths=[2*inch, 4*inch])
    patient_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f8f9fa')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements.append(patient_table)
    elements.append(Spacer(1, 0.3*inch))
    
    # Analysis Results Section
    elements.append(Paragraph("Analysis Results", heading_style))
    
    # Determine result color and status
    if prediction_label == "Pneumonia":
        result_color = colors.HexColor('#dc3545')
        status_text = "Pneumonia Detected"
        recommendation = "The analysis indicates signs consistent with pneumonia. Please consult with a healthcare professional for proper diagnosis and treatment."
    else:
        result_color = colors.HexColor('#28a745')
        status_text = "Normal - No Pneumonia Detected"
        recommendation = "The analysis did not detect signs of pneumonia. However, this is not a substitute for professional medical evaluation."
    
    confidence_pct = confidence * 100
    
    # Results table
    results_data = [
        ['Diagnosis:', status_text],
        ['Confidence Level:', f'{confidence_pct:.1f}%'],
        ['Analysis Status:', 'High Confidence' if confidence > 0.8 else 'Medium Confidence' if confidence > 0.6 else 'Low Confidence']
    ]
    
    results_table = Table(results_data, colWidths=[2*inch, 4*inch])
    results_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f8f9fa')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('TEXTCOLOR', (1, 0), (1, 0), result_color),  # Diagnosis in color
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements.append(results_table)
    elements.append(Spacer(1, 0.2*inch))
    
    # Recommendation
    rec_style = ParagraphStyle(
        'Recommendation',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#495057'),
        spaceAfter=12,
        alignment=TA_LEFT,
        fontName='Helvetica-Oblique',
        backColor=colors.HexColor('#fff3cd'),
        borderPadding=10
    )
    elements.append(Paragraph(f"<b>Recommendation:</b> {recommendation}", rec_style))
    elements.append(Spacer(1, 0.3*inch))
    
    # Images Section
    elements.append(Paragraph("X-Ray Images", heading_style))
    
    # Convert PIL images to bytes for reportlab
    def pil_to_bytes(img, max_width=3*inch):
        """Convert PIL Image to bytes and resize if needed"""
        img_copy = img.copy()
        # Resize if too large
        if img_copy.width > max_width:
            ratio = max_width / img_copy.width
            new_height = int(img_copy.height * ratio)
            img_copy = img_copy.resize((int(max_width), new_height), PILImage.Resampling.LANCZOS)
        
        img_bytes = io.BytesIO()
        img_copy.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        return img_bytes
    
    # Original image
    try:
        orig_img_bytes = pil_to_bytes(original_image, max_width=3*inch)
        orig_img = ReportLabImage(orig_img_bytes, width=3*inch, height=3*inch * (original_image.height / original_image.width))
        elements.append(Paragraph("<b>Original Chest X-Ray</b>", normal_style))
        elements.append(orig_img)
        elements.append(Spacer(1, 0.2*inch))
    except Exception as e:
        elements.append(Paragraph(f"<i>Original image could not be included: {str(e)}</i>", normal_style))
    
    # Heatmap image
    try:
        heatmap_img_bytes = pil_to_bytes(heatmap_image, max_width=3*inch)
        heatmap_img = ReportLabImage(heatmap_img_bytes, width=3*inch, height=3*inch * (heatmap_image.height / heatmap_image.width))
        elements.append(Paragraph("<b>Analysis Heatmap (Key Regions Highlighted)</b>", normal_style))
        elements.append(heatmap_img)
        elements.append(Spacer(1, 0.3*inch))
    except Exception as e:
        elements.append(Paragraph(f"<i>Heatmap image could not be included: {str(e)}</i>", normal_style))
    
    # Medical Disclaimer
    elements.append(Spacer(1, 0.2*inch))
    disclaimer_style = ParagraphStyle(
        'Disclaimer',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#6c757d'),
        spaceAfter=6,
        alignment=TA_LEFT,
        fontName='Helvetica-Oblique',
        backColor=colors.HexColor('#f8f9fa'),
        borderPadding=10
    )
    
    disclaimer_text = """
    <b>IMPORTANT MEDICAL DISCLAIMER:</b><br/>
    This report is generated using automated analysis tools and is provided for educational and research purposes only. 
    This application is NOT a substitute for professional medical diagnosis. Always consult qualified healthcare 
    professionals for medical decisions. Do not use this tool as the sole basis for treatment decisions. 
    Results may contain inaccuracies and should be reviewed and verified by medical experts.
    """
    elements.append(Paragraph(disclaimer_text, disclaimer_style))
    
    # Digital Signature Section
    elements.append(Spacer(1, 0.4*inch))
    elements.append(Paragraph("Digital Signature", heading_style))
    elements.append(Spacer(1, 0.1*inch))
    
    # Create and add signature image
    try:
        signature_img = create_digital_signature()
        signature_bytes = io.BytesIO()
        signature_img.save(signature_bytes, format='PNG')
        signature_bytes.seek(0)
        
        sig_img = ReportLabImage(signature_bytes, width=4*inch, height=1.2*inch)
        elements.append(sig_img)
    except Exception as e:
        # Fallback to text if image creation fails
        signature_style = ParagraphStyle(
            'Signature',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#2c5aa0'),
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        elements.append(Paragraph("Keep Awareness Alive", signature_style))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.HexColor('#6c757d'),
            alignment=TA_CENTER
        )
        elements.append(Paragraph("Medical AI Analysis Platform", footer_style))
    
    elements.append(Spacer(1, 0.2*inch))
    
    # Footer
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.HexColor('#6c757d'),
        alignment=TA_CENTER
    )
    elements.append(Paragraph(f"Report ID: {datetime.now().strftime('%Y%m%d%H%M%S')}", footer_style))
    
    # Build PDF
    doc.build(elements)
    return output_path


def create_excel_report(username, email, prediction_label, confidence, original_image, heatmap_image, output_path,
                        patient_name=None, patient_id=None, date_of_birth=None, gender=None, age=None):
    """
    Create an Excel report for the patient
    
    Args:
        username: Patient/Doctor username
        email: Patient/Doctor email
        prediction_label: "Normal" or "Pneumonia"
        confidence: Confidence score (0-1)
        original_image: PIL Image of original X-ray
        heatmap_image: PIL Image of heatmap overlay
        output_path: Path to save the Excel file
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Analysis Report"
    
    # Styles
    title_font = Font(name='Arial', size=16, bold=True, color='FFFFFF')
    heading_font = Font(name='Arial', size=12, bold=True)
    normal_font = Font(name='Arial', size=11)
    
    header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
    normal_fill = PatternFill(start_color='F2F2F2', end_color='F2F2F2', fill_type='solid')
    pneumonia_fill = PatternFill(start_color='FFE6E6', end_color='FFE6E6', fill_type='solid')
    normal_result_fill = PatternFill(start_color='E6F7E6', end_color='E6F7E6', fill_type='solid')
    
    center_align = Alignment(horizontal='center', vertical='center')
    left_align = Alignment(horizontal='left', vertical='center')
    
    # Title
    ws.merge_cells('A1:B1')
    title_cell = ws['A1']
    title_cell.value = "Chest X-Ray Analysis Report"
    title_cell.font = title_font
    title_cell.fill = header_fill
    title_cell.alignment = center_align
    ws.row_dimensions[1].height = 30
    
    # Patient Information
    ws['A3'] = "Patient Information"
    ws['A3'].font = heading_font
    ws.merge_cells('A3:B3')
    
    patient_data = []
    if patient_name:
        patient_data.append(['Patient Name:', patient_name])
    if patient_id:
        patient_data.append(['Patient ID:', patient_id])
    if date_of_birth:
        patient_data.append(['Date of Birth:', date_of_birth])
    if gender:
        patient_data.append(['Gender:', gender])
    if age:
        patient_data.append(['Age:', str(age)])
    
    patient_data.append(['Doctor Name:', username])
    patient_data.append(['Doctor Email:', email])
    patient_data.append(['Report Date:', datetime.now().strftime("%B %d, %Y at %I:%M %p")])
    
    for idx, (label, value) in enumerate(patient_data, start=4):
        ws[f'A{idx}'] = label
        ws[f'A{idx}'].font = heading_font
        ws[f'A{idx}'].fill = normal_fill
        ws[f'A{idx}'].alignment = left_align
        
        ws[f'B{idx}'] = value
        ws[f'B{idx}'].font = normal_font
        ws[f'B{idx}'].alignment = left_align
    
    # Analysis Results
    ws['A8'] = "Analysis Results"
    ws['A8'].font = heading_font
    ws.merge_cells('A8:B8')
    
    confidence_pct = confidence * 100
    status_text = "Pneumonia Detected" if prediction_label == "Pneumonia" else "Normal - No Pneumonia Detected"
    analysis_status = 'High Confidence' if confidence > 0.8 else 'Medium Confidence' if confidence > 0.6 else 'Low Confidence'
    
    results_data = [
        ['Diagnosis:', status_text],
        ['Confidence Level:', f'{confidence_pct:.1f}%'],
        ['Analysis Status:', analysis_status]
    ]
    
    for idx, (label, value) in enumerate(results_data, start=9):
        ws[f'A{idx}'] = label
        ws[f'A{idx}'].font = heading_font
        ws[f'A{idx}'].fill = normal_fill
        ws[f'A{idx}'].alignment = left_align
        
        result_cell = ws[f'B{idx}']
        result_cell.value = value
        result_cell.font = normal_font
        result_cell.alignment = left_align
        
        if idx == 9:  # Diagnosis row
            if prediction_label == "Pneumonia":
                result_cell.fill = pneumonia_fill
            else:
                result_cell.fill = normal_result_fill
    
    # Recommendation
    ws['A13'] = "Recommendation"
    ws['A13'].font = heading_font
    ws.merge_cells('A13:B13')
    
    recommendation = ("The analysis indicates signs consistent with pneumonia. Please consult with a healthcare "
                     "professional for proper diagnosis and treatment." if prediction_label == "Pneumonia" 
                     else "The analysis did not detect signs of pneumonia. However, this is not a substitute for "
                     "professional medical evaluation.")
    
    ws['A14'] = recommendation
    ws['A14'].font = normal_font
    ws.merge_cells('A14:B14')
    ws['A14'].alignment = left_align
    ws.row_dimensions[14].height = 40
    
    # Image paths note
    ws['A16'] = "Note: Images are saved as separate files"
    ws['A16'].font = Font(name='Arial', size=9, italic=True)
    ws.merge_cells('A16:B16')
    
    # Disclaimer
    ws['A18'] = "IMPORTANT MEDICAL DISCLAIMER:"
    ws['A18'].font = Font(name='Arial', size=10, bold=True)
    ws.merge_cells('A18:B18')
    
    disclaimer = ("This report is generated using automated analysis tools and is provided for educational and "
                 "research purposes only. This application is NOT a substitute for professional medical diagnosis. "
                 "Always consult qualified healthcare professionals for medical decisions.")
    
    ws['A19'] = disclaimer
    ws['A19'].font = Font(name='Arial', size=9, italic=True)
    ws.merge_cells('A19:B19')
    ws['A19'].alignment = left_align
    ws.row_dimensions[19].height = 50
    
    # Footer
    # Digital Signature Section
    sig_row = 21
    ws[f'A{sig_row}'] = "Digital Signature"
    ws[f'A{sig_row}'].font = heading_font
    ws.merge_cells(f'A{sig_row}:B{sig_row}')
    
    sig_row += 1
    ws[f'A{sig_row}'] = "Keep Awareness Alive"
    ws[f'A{sig_row}'].font = Font(name='Arial', size=14, bold=True, color='2c5aa0')
    ws.merge_cells(f'A{sig_row}:B{sig_row}')
    ws[f'A{sig_row}'].alignment = center_align
    
    sig_row += 1
    ws[f'A{sig_row}'] = "Medical AI Analysis Platform"
    ws[f'A{sig_row}'].font = Font(name='Arial', size=10, italic=True, color='666666')
    ws.merge_cells(f'A{sig_row}:B{sig_row}')
    ws[f'A{sig_row}'].alignment = center_align
    
    sig_row += 1
    ws[f'A{sig_row}'] = datetime.now().strftime("%B %d, %Y")
    ws[f'A{sig_row}'].font = Font(name='Arial', size=9, color='666666')
    ws.merge_cells(f'A{sig_row}:B{sig_row}')
    ws[f'A{sig_row}'].alignment = center_align
    
    # Footer
    footer_row = sig_row + 2
    ws[f'A{footer_row}'] = f"Report ID: {datetime.now().strftime('%Y%m%d%H%M%S')}"
    ws[f'A{footer_row}'].font = Font(name='Arial', size=8)
    ws.merge_cells(f'A{footer_row}:B{footer_row}')
    ws[f'A{footer_row}'].alignment = center_align
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 40
    
    # Save workbook
    wb.save(output_path)
    return output_path


def create_word_report(username, email, prediction_label, confidence, original_image, heatmap_image, output_path,
                       patient_name=None, patient_id=None, date_of_birth=None, gender=None, age=None):
    """
    Create a Word document report for the patient
    
    Args:
        username: Patient/Doctor username
        email: Patient/Doctor email
        prediction_label: "Normal" or "Pneumonia"
        confidence: Confidence score (0-1)
        original_image: PIL Image of original X-ray
        heatmap_image: PIL Image of heatmap overlay
        output_path: Path to save the Word document
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('Chest X-Ray Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Patient Information
    doc.add_heading('Patient Information', level=1)
    
    # Build patient data list
    patient_data = []
    if patient_name:
        patient_data.append(['Patient Name:', patient_name])
    if patient_id:
        patient_data.append(['Patient ID:', patient_id])
    if date_of_birth:
        patient_data.append(['Date of Birth:', date_of_birth])
    if gender:
        patient_data.append(['Gender:', gender])
    if age:
        patient_data.append(['Age:', str(age)])
    
    patient_data.append(['Doctor Name:', username])
    patient_data.append(['Doctor Email:', email])
    patient_data.append(['Report Date:', datetime.now().strftime("%B %d, %Y at %I:%M %p")])
    
    patient_table = doc.add_table(rows=len(patient_data), cols=2)
    patient_table.style = 'Light Grid Accent 1'

    for i, (label, value) in enumerate(patient_data):
        patient_table.rows[i].cells[0].text = label
        patient_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        patient_table.rows[i].cells[1].text = value
    
    # Analysis Results
    doc.add_heading('Analysis Results', level=1)
    
    confidence_pct = confidence * 100
    status_text = "Pneumonia Detected" if prediction_label == "Pneumonia" else "Normal - No Pneumonia Detected"
    analysis_status = 'High Confidence' if confidence > 0.8 else 'Medium Confidence' if confidence > 0.6 else 'Low Confidence'
    
    results_table = doc.add_table(rows=3, cols=2)
    results_table.style = 'Light Grid Accent 1'
    
    results_data = [
        ['Diagnosis:', status_text],
        ['Confidence Level:', f'{confidence_pct:.1f}%'],
        ['Analysis Status:', analysis_status]
    ]
    
    for i, (label, value) in enumerate(results_data):
        results_table.rows[i].cells[0].text = label
        results_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        
        value_cell = results_table.rows[i].cells[1]
        value_cell.text = value
        if i == 0:  # Diagnosis row
            if prediction_label == "Pneumonia":
                value_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 53, 69)
            else:
                value_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(40, 167, 69)
    
    # Recommendation
    doc.add_heading('Recommendation', level=1)
    recommendation = ("The analysis indicates signs consistent with pneumonia. Please consult with a healthcare "
                     "professional for proper diagnosis and treatment." if prediction_label == "Pneumonia" 
                     else "The analysis did not detect signs of pneumonia. However, this is not a substitute for "
                     "professional medical evaluation.")
    
    rec_para = doc.add_paragraph(recommendation)
    rec_para.style = 'Intense Quote'
    
    # Images Section
    doc.add_heading('X-Ray Images', level=1)
    
    # Save images temporarily and add to document
    try:
        # Original image
        orig_img_bytes = io.BytesIO()
        original_image.save(orig_img_bytes, format='PNG')
        orig_img_bytes.seek(0)
        
        doc.add_heading('Original Chest X-Ray', level=2)
        orig_para = doc.add_paragraph()
        orig_run = orig_para.add_run()
        orig_run.add_picture(orig_img_bytes, width=Inches(4))
        
        # Heatmap image
        heatmap_img_bytes = io.BytesIO()
        heatmap_image.save(heatmap_img_bytes, format='PNG')
        heatmap_img_bytes.seek(0)
        
        doc.add_heading('Analysis Heatmap (Key Regions Highlighted)', level=2)
        heatmap_para = doc.add_paragraph()
        heatmap_run = heatmap_para.add_run()
        heatmap_run.add_picture(heatmap_img_bytes, width=Inches(4))
    except Exception as e:
        doc.add_paragraph(f"Note: Images could not be included: {str(e)}")
    
    # Medical Disclaimer
    doc.add_heading('Important Medical Disclaimer', level=1)
    disclaimer_text = ("This report is generated using automated analysis tools and is provided for educational and "
                      "research purposes only. This application is NOT a substitute for professional medical diagnosis. "
                      "Always consult qualified healthcare professionals for medical decisions. Do not use this tool as "
                      "the sole basis for treatment decisions. Results may contain inaccuracies and should be reviewed "
                      "and verified by medical experts.")
    
    disclaimer_para = doc.add_paragraph(disclaimer_text)
    disclaimer_para.style = 'Intense Quote'
    for run in disclaimer_para.runs:
        run.font.italic = True
        run.font.size = Pt(9)
    
    # Digital Signature Section
    doc.add_paragraph()
    doc.add_heading('Digital Signature', level=1)
    
    # Create and add signature
    try:
        signature_img = create_digital_signature()
        signature_bytes = io.BytesIO()
        signature_img.save(signature_bytes, format='PNG')
        signature_bytes.seek(0)
        
        sig_para = doc.add_paragraph()
        sig_run = sig_para.add_run()
        sig_run.add_picture(signature_bytes, width=Inches(4))
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        # Fallback to text
        sig_para = doc.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig_run = sig_para.add_run("Keep Awareness Alive")
        sig_run.font.size = Pt(14)
        sig_run.font.bold = True
        sig_run.font.color.rgb = RGBColor(44, 90, 160)
        
        tagline_para = doc.add_paragraph()
        tagline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline_run = tagline_para.add_run("Medical AI Analysis Platform")
        tagline_run.font.size = Pt(10)
        tagline_run.font.italic = True
        tagline_run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph(f"Report ID: {datetime.now().strftime('%Y%m%d%H%M%S')}")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(8)
    footer_para.runs[0].font.color.rgb = RGBColor(108, 117, 125)
    
    doc.save(output_path)
    return output_path


def create_csv_report(username, email, prediction_label, confidence, original_image, heatmap_image, output_path,
                      patient_name=None, patient_id=None, date_of_birth=None, gender=None, age=None):
    """
    Create a CSV report for the patient
    
    Args:
        username: Patient/Doctor username
        email: Patient/Doctor email
        prediction_label: "Normal" or "Pneumonia"
        confidence: Confidence score (0-1)
        original_image: PIL Image of original X-ray
        heatmap_image: PIL Image of heatmap overlay
        output_path: Path to save the CSV file
    """
    confidence_pct = confidence * 100
    status_text = "Pneumonia Detected" if prediction_label == "Pneumonia" else "Normal - No Pneumonia Detected"
    analysis_status = 'High Confidence' if confidence > 0.8 else 'Medium Confidence' if confidence > 0.6 else 'Low Confidence'
    
    fields = ['Report Title']
    values = ['Chest X-Ray Analysis Report']
    
    if patient_name:
        fields.append('Patient Name')
        values.append(patient_name)
    if patient_id:
        fields.append('Patient ID')
        values.append(patient_id)
    if date_of_birth:
        fields.append('Date of Birth')
        values.append(date_of_birth)
    if gender:
        fields.append('Gender')
        values.append(gender)
    if age:
        fields.append('Age')
        values.append(str(age))
    
    fields.extend(['Doctor Name', 'Doctor Email', 'Report Date', 'Diagnosis', 'Confidence Level (%)', 
                   'Analysis Status', 'Recommendation', 'Report ID', 'Disclaimer'])
    values.extend([
        username,
        email,
        datetime.now().strftime("%B %d, %Y at %I:%M %p"),
        status_text,
        f'{confidence_pct:.1f}',
        analysis_status,
        ("The analysis indicates signs consistent with pneumonia. Please consult with a healthcare "
         "professional for proper diagnosis and treatment." if prediction_label == "Pneumonia" 
         else "The analysis did not detect signs of pneumonia. However, this is not a substitute for "
         "professional medical evaluation."),
        datetime.now().strftime('%Y%m%d%H%M%S'),
        ("This report is generated using automated analysis tools and is provided for educational and "
         "research purposes only. This application is NOT a substitute for professional medical diagnosis.")
    ])
    
    data = {
        'Field': fields,
        'Value': values
    }
    
    df = pd.DataFrame(data)
    df.to_csv(output_path, index=False)
    return output_path


def create_json_report(username, email, prediction_label, confidence, original_image, heatmap_image, output_path,
                       patient_name=None, patient_id=None, date_of_birth=None, gender=None, age=None):
    """
    Create a JSON report for the patient
    
    Args:
        username: Patient/Doctor username
        email: Patient/Doctor email
        prediction_label: "Normal" or "Pneumonia"
        confidence: Confidence score (0-1)
        original_image: PIL Image of original X-ray
        heatmap_image: PIL Image of heatmap overlay
        output_path: Path to save the JSON file
    """
    confidence_pct = confidence * 100
    status_text = "Pneumonia Detected" if prediction_label == "Pneumonia" else "Normal - No Pneumonia Detected"
    analysis_status = 'High Confidence' if confidence > 0.8 else 'Medium Confidence' if confidence > 0.6 else 'Low Confidence'
    
    # Convert images to base64
    def image_to_base64(img):
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        return base64.b64encode(img_bytes.read()).decode('utf-8')
    
    patient_info = {}
    if patient_name:
        patient_info['patient_name'] = patient_name
    if patient_id:
        patient_info['patient_id'] = patient_id
    if date_of_birth:
        patient_info['date_of_birth'] = date_of_birth
    if gender:
        patient_info['gender'] = gender
    if age:
        patient_info['age'] = age
    
    report_data = {
        'report_title': 'Chest X-Ray Analysis Report',
        'patient_information': {
            **patient_info,
            'doctor_name': username,
            'doctor_email': email,
            'report_date': datetime.now().strftime("%B %d, %Y at %I:%M %p"),
            'report_id': datetime.now().strftime('%Y%m%d%H%M%S')
        },
        'analysis_results': {
            'diagnosis': status_text,
            'prediction_label': prediction_label,
            'confidence_level': confidence_pct,
            'confidence_raw': confidence,
            'analysis_status': analysis_status
        },
        'recommendation': ("The analysis indicates signs consistent with pneumonia. Please consult with a healthcare "
                          "professional for proper diagnosis and treatment." if prediction_label == "Pneumonia" 
                          else "The analysis did not detect signs of pneumonia. However, this is not a substitute for "
                          "professional medical evaluation."),
        'images': {
            'original_image_base64': image_to_base64(original_image),
            'heatmap_image_base64': image_to_base64(heatmap_image)
        },
        'disclaimer': ("This report is generated using automated analysis tools and is provided for educational and "
                      "research purposes only. This application is NOT a substitute for professional medical diagnosis. "
                      "Always consult qualified healthcare professionals for medical decisions."),
        'digital_signature': {
            'organization': 'Keep Awareness Alive',
            'tagline': 'Medical AI Analysis Platform',
            'signature_date': datetime.now().strftime("%B %d, %Y")
        },
        'metadata': {
            'generated_by': 'Keep Awareness Alive',
            'generation_timestamp': datetime.now().isoformat()
        }
    }
    
    with open(output_path, 'w') as f:
        json.dump(report_data, f, indent=4)
    
    return output_path


def create_html_report(username, email, prediction_label, confidence, original_image, heatmap_image, output_path,
                      patient_name=None, patient_id=None, date_of_birth=None, gender=None, age=None):
    """
    Create an HTML report for the patient
    
    Args:
        username: Patient/Doctor username
        email: Patient/Doctor email
        prediction_label: "Normal" or "Pneumonia"
        confidence: Confidence score (0-1)
        original_image: PIL Image of original X-ray
        heatmap_image: PIL Image of heatmap overlay
        output_path: Path to save the HTML file
    """
    confidence_pct = confidence * 100
    status_text = "Pneumonia Detected" if prediction_label == "Pneumonia" else "Normal - No Pneumonia Detected"
    analysis_status = 'High Confidence' if confidence > 0.8 else 'Medium Confidence' if confidence > 0.6 else 'Low Confidence'
    
    # Convert images to base64
    def image_to_base64(img):
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        return base64.b64encode(img_bytes.read()).decode('utf-8')
    
    result_color = '#dc3545' if prediction_label == "Pneumonia" else '#28a745'
    recommendation = ("The analysis indicates signs consistent with pneumonia. Please consult with a healthcare "
                     "professional for proper diagnosis and treatment." if prediction_label == "Pneumonia" 
                     else "The analysis did not detect signs of pneumonia. However, this is not a substitute for "
                     "professional medical evaluation.")
    
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Chest X-Ray Analysis Report</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                max-width: 900px;
                margin: 0 auto;
                padding: 20px;
                background-color: #f5f5f5;
            }}
            .container {{
                background-color: white;
                padding: 30px;
                border-radius: 10px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            }}
            h1 {{
                color: #1a1a1a;
                text-align: center;
                border-bottom: 3px solid #366092;
                padding-bottom: 10px;
            }}
            h2 {{
                color: #2c3e50;
                margin-top: 30px;
                border-bottom: 2px solid #e0e0e0;
                padding-bottom: 5px;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
            }}
            th, td {{
                padding: 12px;
                text-align: left;
                border-bottom: 1px solid #ddd;
            }}
            th {{
                background-color: #f8f9fa;
                font-weight: bold;
                width: 30%;
            }}
            .diagnosis {{
                color: {result_color};
                font-weight: bold;
                font-size: 1.1em;
            }}
            .recommendation {{
                background-color: #fff3cd;
                padding: 15px;
                border-left: 4px solid #ffc107;
                margin: 20px 0;
                font-style: italic;
            }}
            .disclaimer {{
                background-color: #f8f9fa;
                padding: 15px;
                border-left: 4px solid #6c757d;
                margin: 20px 0;
                font-size: 0.9em;
                color: #6c757d;
            }}
            .image-container {{
                text-align: center;
                margin: 20px 0;
            }}
            .image-container img {{
                max-width: 100%;
                height: auto;
                border: 2px solid #ddd;
                border-radius: 5px;
                margin: 10px;
            }}
            .footer {{
                text-align: center;
                margin-top: 30px;
                padding-top: 20px;
                border-top: 1px solid #ddd;
                color: #6c757d;
                font-size: 0.9em;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1>Chest X-Ray Analysis Report</h1>
            
            <h2>Patient Information</h2>
            <table>
                {f'<tr><th>Patient Name:</th><td>{patient_name}</td></tr>' if patient_name else ''}
                {f'<tr><th>Patient ID:</th><td>{patient_id}</td></tr>' if patient_id else ''}
                {f'<tr><th>Date of Birth:</th><td>{date_of_birth}</td></tr>' if date_of_birth else ''}
                {f'<tr><th>Gender:</th><td>{gender}</td></tr>' if gender else ''}
                {f'<tr><th>Age:</th><td>{age}</td></tr>' if age else ''}
                <tr>
                    <th>Doctor Name:</th>
                    <td>{username}</td>
                </tr>
                <tr>
                    <th>Doctor Email:</th>
                    <td>{email}</td>
                </tr>
                <tr>
                    <th>Report Date:</th>
                    <td>{datetime.now().strftime("%B %d, %Y at %I:%M %p")}</td>
                </tr>
            </table>
            
            <h2>Analysis Results</h2>
            <table>
                <tr>
                    <th>Diagnosis:</th>
                    <td class="diagnosis">{status_text}</td>
                </tr>
                <tr>
                    <th>Confidence Level:</th>
                    <td>{confidence_pct:.1f}%</td>
                </tr>
                <tr>
                    <th>Analysis Status:</th>
                    <td>{analysis_status}</td>
                </tr>
            </table>
            
            <div class="recommendation">
                <strong>Recommendation:</strong> {recommendation}
            </div>
            
            <h2>X-Ray Images</h2>
            <div class="image-container">
                <h3>Original Chest X-Ray</h3>
                <img src="data:image/png;base64,{image_to_base64(original_image)}" alt="Original X-Ray">
                
                <h3>Analysis Heatmap (Key Regions Highlighted)</h3>
                <img src="data:image/png;base64,{image_to_base64(heatmap_image)}" alt="Heatmap Overlay">
            </div>
            
            <div class="disclaimer">
                <strong>IMPORTANT MEDICAL DISCLAIMER:</strong><br>
                This report is generated using automated analysis tools and is provided for educational and research purposes only. 
                This application is NOT a substitute for professional medical diagnosis. Always consult qualified healthcare 
                professionals for medical decisions. Do not use this tool as the sole basis for treatment decisions. 
                Results may contain inaccuracies and should be reviewed and verified by medical experts.
            </div>
            
            <div style="margin-top: 40px; padding-top: 20px; border-top: 2px solid #e0e0e0;">
                <h3 style="text-align: center; color: #2c5aa0; margin-bottom: 10px;">Digital Signature</h3>
                <div style="text-align: center; margin: 20px 0;">
                    <div style="display: inline-block; text-align: center; padding: 20px; border: 1px solid #e0e0e0; border-radius: 5px; background-color: #f9f9f9; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                        <div style="font-size: 24px; font-weight: bold; color: #2c5aa0; margin-bottom: 5px;">Keep Awareness Alive</div>
                        <div style="font-size: 14px; color: #666; font-style: italic; margin-bottom: 10px;">Medical AI Analysis Platform</div>
                        <div style="font-size: 12px; color: #999; margin-top: 10px; padding-top: 10px; border-top: 1px solid #e0e0e0;">{datetime.now().strftime("%B %d, %Y")}</div>
                    </div>
                </div>
            </div>
            
            <div class="footer">
                Report ID: {datetime.now().strftime('%Y%m%d%H%M%S')}
            </div>
        </div>
    </body>
    </html>
    """
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    return output_path

